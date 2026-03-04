import os
import random
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importar el nuevo orquestador de narrativas
try:
    from app.generador_narrador import generar_narrativa as generar_narrativa_orquestador
except ImportError:
    from generador_narrador import generar_narrativa as generar_narrativa_orquestador

EXPORTS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../exports'))

DIMENSIONES = [
    {
        "nombre": "Dimensión Geográfica",
        "descripcion": "El RSEIA define esta dimensión como la distribución de los grupos humanos en el territorio y sus sistemas de movilidad y conectividad.",
        "variables": ["Medio de transporte principal al trabajo"]
    },
    {
        "nombre": "Dimensión Demográfica",
        "descripcion": "Corresponde a la estructura de la población local, tamaño, densidad, y características intrínsecas de los habitantes.",
        "variables": ["Población Total", "Edad en tramos", "Estado conyugal o civil", "Discapacidad"]
    },
    {
        "nombre": "Dimensión Antropológica",
        "descripcion": "Abarca las características étnicas, manifestaciones de la cultura, sistemas de valores y creencias.",
        "variables": ["Autoiden. pueblo indígena u originario", "A qué pueblo indígena pertenece", "Afrodescendencia", "Tiene religión o credo"]
    },
    {
        "nombre": "Dimensión Socioeconómica",
        "descripcion": "Agrupa los factores relativos al empleo, niveles de ingreso, actividad económica predominante y el capital humano o patrimonial.",
        "variables": ["Tenencia de vivienda", "Años de escolaridad", "Situación en la fuerza de trabajo", "Categoría ocupacional"]
    },
    {
        "nombre": "Dimensión Bienestar Social Básico",
        "descripcion": "El RSEIA es claro en que esta dimensión mide el acceso a la vivienda, equipamiento y servicios básicos (agua, energía, saneamiento, comunicaciones). Esta es la dimensión más robusta en tu listado, ya que gran parte de los datos provienen de la caracterización material del hogar.",
        "variables": ["Tipo de vivienda particular", "Estado de ocupación de la vivienda", "Material de construcción en las paredes exteriores", "Material de construcción en la cubierta del techo", "Material de construcción en el piso", "Origen del agua (Fuente principal)", "Sistema de distribución del agua", "Servicio higiénico (WC)", "Origen de la electricidad", "Medio de eliminación de basura", "Combustible para cocinar", "Combustible para calefaccionar", "Dispone de internet fija"]
    }
]

def add_seq_field(run, seq_identifier):
    """Inserta un campo SEQ en un run de docx para numeración automática."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' SEQ {seq_identifier} \\* ARABIC '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    t = OxmlElement('w:t')
    t.text = "1"

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

os.makedirs(EXPORTS_DIR, exist_ok=True)

def generar_narrativa(titulo: str, datos: dict, num_tabla: int) -> str:
    """Genera un párrafo descriptivo automático basado en los datos agrupados, invocando al nuevo orquestador."""
    unidad = datos.get("Unidad", "registros")
    total = datos.get("Denominador", 0)
    cats = datos.get("Categorias", {})
    
    if total == 0 or not cats:
        return f"No hay datos suficientes para generar un resumen para {titulo}."

    # Mantenemos a salvo la lógica excepcional de Población Total que no opera como las demás
    if "Poblacion Total" in titulo or "Población Total" in titulo:
        total_data = cats.get("Total", {})
        hombres = total_data.get("Hombres", 0)
        mujeres = total_data.get("Mujeres", 0)
        pct_h = round((hombres / total * 100), 1) if total > 0 else 0
        pct_m = round((mujeres / total * 100), 1) if total > 0 else 0
        
        return f"La población total de la comuna es de {total:,} {unidad}. De estos, un {pct_m}% corresponden a mujeres y un {pct_h}% a hombres.\n\nEl siguiente cuadro presenta la tabla con información detallada."

    # Procesar PctMujeresTop1 para enviarlo al orquestador
    # 1. Encontrar la categoría más frecuente (excluyendo "No respuesta", "No declarado", "Ignorado", "Total")
    cats_validas = {k: v for k, v in cats.items() if k not in ["No respuesta", "No declarado", "Ignorado", "Total"] and "Subtotal" in v}
    pct_mujeres_top1 = None
    
    if cats_validas:
        cat_top = max(cats_validas.keys(), key=lambda k: cats_validas[k].get("Subtotal", 0))
        top_data = cats_validas[cat_top]
        if "Mujeres" in top_data and "Hombres" in top_data:
            sum_sexo = top_data["Mujeres"] + top_data["Hombres"]
            if sum_sexo > 0:
                pct_mujeres_top1 = round((top_data["Mujeres"] / sum_sexo) * 100, 1)
                
    # Inyectamos el valor calculado al diccionario in-situ para el orquestador
    datos["PctMujeresTop1"] = pct_mujeres_top1

    # Invocamos el orquestador principal
    return generar_narrativa_orquestador(titulo, datos, num_tabla)

def _crear_dataframe_desde_diccionario(datos: dict) -> pd.DataFrame:
    """Convierte el diccionario de Categorias a un Pandas DataFrame plano para exportar."""
    filas = []
    cats = datos.get("Categorias", {})
    
    # Check if "Total" is the only key (e.g., Poblacion Total)
    is_only_total = (list(cats.keys()) == ["Total"])
    
    for k, v in cats.items():
        if k == "Total" and not is_only_total:
            continue # Lo manejamos al final si está presente
        fila = {"Categoría": k}
        fila.update(v)
        
        # Calcular porcentajes por sexo si están presentes
        if "Subtotal" in v and "Hombres" in v and "Mujeres" in v:
            frecuencia = v["Subtotal"]
            if frecuencia > 0:
                fila["Hombres (%)"] = round((v["Hombres"] / frecuencia) * 100, 1)
                fila["Mujeres (%)"] = round((v["Mujeres"] / frecuencia) * 100, 1)
            else:
                fila["Hombres (%)"] = 0.0
                fila["Mujeres (%)"] = 0.0

        filas.append(fila)
        
    df = pd.DataFrame(filas)
    
    # Renombrar Subtotal a Frecuencia
    if 'Subtotal' in df.columns:
        df.rename(columns={'Subtotal': 'Frecuencia'}, inplace=True)
        
    # Calcular y agregar fila Total (excluyendo Porcentaje temporalmente)
    # Solo agregamos un nuevo "Total" si no existía ya explícitamente como única fila
    if not df.empty and not is_only_total:
        total_row = {"Categoría": "Total"}
        for col in df.columns:
            if col not in ["Categoría", "Porcentaje", "Hombres (%)", "Mujeres (%)"]:
                total_row[col] = df[col].sum()
        
        # El Porcentaje total siempre será la suma de los porcentajes o 100 si todos informaron
        if "Porcentaje" in df.columns:
            total_row["Porcentaje"] = round(df["Porcentaje"].sum(), 1)
            
        # Calcular porcentajes de sexo para la fila total
        if "Frecuencia" in total_row and "Hombres" in total_row and "Mujeres" in total_row:
            total_frecuencia = total_row["Frecuencia"]
            if total_frecuencia > 0:
                total_row["Hombres (%)"] = round((total_row["Hombres"] / total_frecuencia) * 100, 1)
                total_row["Mujeres (%)"] = round((total_row["Mujeres"] / total_frecuencia) * 100, 1)
            else:
                total_row["Hombres (%)"] = 0.0
                total_row["Mujeres (%)"] = 0.0

        df = pd.concat([df, pd.DataFrame([total_row])], ignore_index=True)
        
    # Formatear porcentajes con % (e.g. 25.0%)
    for pct_col in ["Porcentaje", "Hombres (%)", "Mujeres (%)"]:
        if pct_col in df.columns:
            df[pct_col] = df[pct_col].apply(lambda x: f"{float(x):.1f}%" if pd.notnull(x) and str(x).strip() != "" else "")
            
    # Reordenar y limpiar columnas según requerimiento del usuario (Frecuencia, Mujeres %, Hombres %, Porcentaje)
    target_cols = ["Categoría", "Frecuencia"]
    if "Mujeres (%)" in df.columns and "Hombres (%)" in df.columns:
        target_cols.extend(["Mujeres (%)", "Hombres (%)"])
    if "Porcentaje" in df.columns:
        target_cols.append("Porcentaje")
        
    for col in df.columns:
        if col not in target_cols and col not in ["Hombres", "Mujeres"]:
            target_cols.append(col)
            
    final_cols = [c for c in target_cols if c in df.columns]
    df = df[final_cols]
            
    return df

def generar_excel(comuna_nombre: str, region_nombre: str, datos_censo: dict, timestamp_str: str) -> str:
    filename = f"Reporte_Censo_{comuna_nombre}_{region_nombre}_{timestamp_str}.xlsx".replace(" ", "_")
    filepath = os.path.join(EXPORTS_DIR, filename)

    with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
        for titulo, info in datos_censo.items():
            df = _crear_dataframe_desde_diccionario(info)
            # Limitar tamaño de nombre de pestaña a 31 chars max por requerimiento de Excel
            sheet_name = titulo[:31].replace(":", "").replace("/", "-")
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            
    return filepath

def generar_word(comuna_nombre: str, region_nombre: str, datos_censo: dict, timestamp_str: str) -> str:
    filename = f"Reporte_Censo_{comuna_nombre}_{region_nombre}_{timestamp_str}.docx".replace(" ", "_")
    filepath = os.path.join(EXPORTS_DIR, filename)

    doc = Document()
    
    # Forzar actualización de campos (SEQ) al abrir el documento en Word
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    doc.settings.element.append(update_fields)
    
    doc.add_heading(f'Reporte Demográfico y Censo 2024', 0)
    # Use level 2 for region/comuna subtitle so the variables can be level 1
    doc.add_heading(f'Región: {region_nombre} | Comuna: {comuna_nombre}', 2)
    doc.add_paragraph(f'Generado el: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_page_break()
    
    # Contador global de tablas del documento
    contador_tabla_global = 1

    for dim in DIMENSIONES:
        dim_vars = {k: v for k, v in datos_censo.items() if k in dim["variables"]}
        
        # Omitir dimensión si no vino ninguna de sus variables seleccionadas
        if not dim_vars:
            continue
            
        # 1. Título de nivel 1 para la Dimensión con numeración SEQ
        p_dim = doc.add_heading('', level=1)
        run_dim = p_dim.add_run()
        add_seq_field(run_dim, "Titulo")
        p_dim.add_run(f". {dim['nombre']}")
        
        # Párrafo descriptivo de la dimensión
        doc.add_paragraph(dim["descripcion"])
        doc.add_paragraph("") # Espacio
        
        # Índice interno para las variables (1.1, 1.2, etc.) visualmente simulado aquí
        var_indice = 1
        
        for titulo, info in dim_vars.items():
            unidad = info.get("Unidad", "registros")
            
            # Título de nivel 2 para la variable 
            # Como Word no enlaza automáticamente los SEQ, usaremos Nivel 2 pero omitimos SEQ de título aquí
            # para no complicar el estilo o simplemente dejamos el Nivel 2 natural de Word.
            doc.add_heading(titulo, level=2)
            
            # Narrativa descriptiva ANTES de la tabla y referenciando a la tabla correcta
            narrativa = generar_narrativa(titulo, info, num_tabla=contador_tabla_global)
            doc.add_paragraph(narrativa)
            
            # 2. Título de la tabla centrado con identificador SEQ, variable, unidad y comuna
            titulo_tabla = titulo.replace("'", "").replace('"', "")
            p_tabla = doc.add_paragraph()
            run_tabla = p_tabla.add_run("Tabla-")
            add_seq_field(run_tabla, "Tabla")
            p_tabla.add_run(f". {titulo_tabla} - {unidad} - Comuna de {comuna_nombre}")
            p_tabla.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            df = _crear_dataframe_desde_diccionario(info)
            
            # Add table
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Light Grid Accent 1'
            
            # Headers
            hdr_cells = table.rows[0].cells
            for col_idx, column in enumerate(df.columns):
                hdr_cells[col_idx].text = str(column)
                
            # Rows
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, val in enumerate(row):
                    if pd.isna(val):
                        row_cells[col_idx].text = ""
                    elif isinstance(val, float):
                        row_cells[col_idx].text = f"{val:.2f}"
                    else:
                        row_cells[col_idx].text = str(val)
                        
            # 3. Footer de la tabla
            p_fuente = doc.add_paragraph("Fuente: Censo, 2024.")
            p_fuente.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_fuente.runs[0]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            doc.add_paragraph("") # Espaciado extra
            var_indice += 1
            contador_tabla_global += 1
        
    doc.save(filepath)
    return filepath
